# -*- coding: utf-8 -*-
import streamlit as st
import os, io, zipfile, tempfile, re, datetime, shutil, random
from pathlib import Path
from docx import Document
from openpyxl import load_workbook

def rub_to_words(n):
    if n == 0: return "ноль"
    units = ["","один","два","три","четыре","пять","шесть","семь","восемь","девять",
             "десять","одиннадцать","двенадцать","тринадцать","четырнадцать","пятнадцать",
             "шестнадцать","семнадцать","восемнадцать","девятнадцать"]
    tens = ["","","двадцать","тридцать","сорок","пятьдесят","шестьдесят","семьдесят","восемьдесят","девяносто"]
    hundreds = ["","сто","двести","триста","четыреста","пятьсот","шестьсот","семьсот","восемьсот","девятьсот"]
    uf = ["","одна","две","три","четыре","пять","шесть","семь","восемь","девять"]
    def _t(n,f=False):
        if n==0: return ""
        r=[]; h=n//100
        if h: r.append(hundreds[h])
        rem=n%100
        if rem<20:
            if rem>0: r.append(uf[rem] if (f and rem in(1,2)) else units[rem])
        else:
            r.append(tens[rem//10])
            if rem%10>0: r.append(uf[rem%10] if (f and rem%10 in(1,2)) else units[rem%10])
        return " ".join(r)
    def _pl(n,a,b,c):
        if 11<=n%100<=19: return c
        l=n%10
        if l==1: return a
        if 2<=l<=4: return b
        return c
    th=n//1000; rm=n%1000; p=[]
    if th>0:
        t=_t(th,True)
        if t: p.append(t+" "+_pl(th,"тысяча","тысячи","тысяч"))
    if rm>0 or not p:
        t=_t(rm)
        if t: p.append(t)
    return " ".join(p)

def date_to_russian(s):
    m=["","января","февраля","марта","апреля","мая","июня",
       "июля","августа","сентября","октября","ноября","декабря"]
    p=s.strip().split(".")
    if len(p)!=3: return s
    return f"{int(p[0])} {m[int(p[1])]} {p[2]} г."

TMPL_DIR = Path(__file__).parent / "templates"

# ============================================================
# Name Banks
# ============================================================
# Russian surnames (60) and given names (55) → 3300 combos
RU_SURNAMES = [
    "Иванов","Петров","Сидоров","Смирнов","Кузнецов","Попов",
    "Васильев","Михайлов","Новиков","Федоров","Морозов","Волков",
    "Алексеев","Лебедев","Семенов","Егоров","Павлов","Козлов",
    "Степанов","Николаев","Орлов","Андреев","Макаров","Никитин",
    "Захаров","Зайцев","Соловьев","Борисов","Яковлев","Григорьев",
    "Романов","Воробьев","Сергеев","Кузьмин","Фролов","Александров",
    "Дмитриев","Королев","Гусев","Киселев","Ильин","Максимов",
    "Поляков","Сорокин","Виноградов","Ковалев","Белов","Медведев",
    "Антонов","Тарасов","Жуков","Баранов","Филиппов","Комаров",
    "Давыдов","Беляев","Герасимов","Богданов","Осипов","Тимофеев"
]
RU_GIVEN = [
    "Александр","Дмитрий","Сергей","Андрей","Алексей","Михаил",
    "Николай","Владимир","Иван","Павел","Роман","Виктор",
    "Юрий","Денис","Евгений","Олег","Игорь","Анатолий",
    "Вадим","Константин","Максим","Антон","Василий","Борис",
    "Геннадий","Григорий","Даниил","Егор","Илья","Кирилл",
    "Лев","Леонид","Матвей","Никита","Петр","Руслан",
    "Святослав","Семен","Станислав","Степан","Тимур","Федор",
    "Эдуард","Ярослав","Артем","Валерий","Владислав","Георгий",
    "Арсений","Валентин","Вячеслав","Тимофей","Всеволод","Марк","Филипп"
]

# Chinese surnames in Russian transliteration (60)
CN_SURNAMES_RU = [
    "Чжан","Ван","Ли","Чжао","Лю","Чэнь","Ян","Хуан","Чжоу","У",
    "Сюй","Сунь","Ма","Чжу","Ху","Го","Хэ","Гао","Линь","Ло",
    "Чжэн","Лян","Се","Сун","Тан","Сюй","Хань","Фэн","Дэн","Цао",
    "Пэн","Цзэн","Сяо","Тянь","Дун","Пань","Юань","Юй","Цзян","Цай",
    "Юй","Ду","Е","Чэн","Су","Вэй","Люй","Дин","Жэнь","Шэнь",
    "Яо","Лу","Цзян","Цуй","Чжун","Тань","Лу","Ван","Фань","Ши"
]
# Chinese given names in Russian transliteration (55)
CN_GIVEN_RU = [
    "Вэй","Лэй","Ян","Юн","Цзюнь","Цзе","Тао","Мин","Чао","Цян",
    "Пэн","Цзяньхуа","Годун","Чжицян","Цзяньго","Вэньбо","Хаожань","Цзыхань",
    "Юйсюань","Юйцзэ","Чжиюань","Сыюань","Бовэнь","Цзюньцзе","Жуй","Чэнь",
    "Ян","Фэн","Нин","Лун","Фэй","Бо","Бинь","Ган","Хуэй","Линь",
    "Минь","Пин","Лян","Синь","И","Сюй","Хао","Сян","Чжэ","Хэн",
    "Юэ","Жань","И","Хань","Цзэ","Жуй","Ань","Кан","Цзянь"
]

# Russian business positions (always used regardless of company type)
RU_POSITIONS = [
    "Генеральный директор","Исполнительный директор","Финансовый директор",
    "Главный бухгалтер","Руководитель отдела продаж","Руководитель транспортного отдела",
    "Специалист по таможенному оформлению","Менеджер по маркетингу","Менеджер по закупкам",
    "Главный инженер","Инженер-проектировщик","Специалист по логистике",
    "Юрисконсульт","Экономист","Специалист по ВЭД","IT-специалист",
    "Начальник склада","Водитель-экспедитор","Переводчик (китайский язык)",
    "Офис-менеджер","Ассистент проекта","Аналитик","Технический директор",
    "Менеджер по продукту","Специалист по безопасности"
]

def generate_random_name(company_type):
    if company_type == "russian":
        return random.choice(RU_SURNAMES) + " " + random.choice(RU_GIVEN)
    else:
        # Chinese name in Russian transliteration: SURNAME + GIVEN (Cyrillic)
        return random.choice(CN_SURNAMES_RU) + " " + random.choice(CN_GIVEN_RU)

def generate_delegation(company_type, count):
    shuffled = list(RU_POSITIONS)
    random.shuffle(shuffled)
    delegation = []
    for i in range(count):
        pos = shuffled[i % len(shuffled)]
        name = generate_random_name(company_type)
        delegation.append((pos, name))
    return delegation

# ============================================================
def get_para_full(para):
    return "".join(r.text for r in para.runs)
def replace_para_text(para, txt):
    for i,r in enumerate(para.runs): r.text = txt if i==0 else ""
def set_cell_text(cell, txt):
    ps=cell.paragraphs
    if ps:
        rs=ps[0].runs
        if rs:
            rs[0].text=str(txt)
            for r in rs[1:]: r.text=""
        else: ps[0].text=str(txt)

def generate_files(data, out_dir):
    date_str = data["date"]
    actual_amt = data["actual_amt"]
    budget_amt = data.get("budget_amt") or actual_amt
    amount_manual = data.get("amount_manual", "").strip()
    company_full = data["company_full"]
    responsible = data["responsible"]
    position = data["position"]
    participant = data["participant"]
    venue = data.get("venue", "")
    address = data.get("address", "")
    purpose = data.get("purpose", "")
    result = data.get("result", "")
    receipt_date = data.get("receipt_date", "")
    receipt_num = data.get("receipt_num", "")
    receipt_amt = data.get("receipt_amt") or actual_amt
    company_short = data.get("company_short") or company_full
    comm1_pos = data.get("comm1_pos", "")
    comm1_name = data.get("comm1_name", "")
    comm2_pos = data.get("comm2_pos", "")
    comm2_name = data.get("comm2_name", "")
    comm3_pos = data.get("comm3_pos", "")
    comm3_name = data.get("comm3_name", "")
    compiler_pos = data.get("compiler_pos", "")
    compiler_name = data.get("compiler_name", "")

    budget_words = amount_manual if amount_manual else rub_to_words(budget_amt)
    actual_words = rub_to_words(actual_amt)
    date_russian = date_to_russian(date_str)
    day, month, year = date_str.split(".")

    rec_date_clean = receipt_date
    for fmt in ["%Y-%m-%d %H:%M:%S", "%Y-%m-%d"]:
        try:
            dt = datetime.datetime.strptime(receipt_date, fmt)
            rec_date_clean = dt.strftime("%d.%m.%Y")
            break
        except: pass

    out_s = os.path.join(out_dir, "Смета.docx")
    out_p = os.path.join(out_dir, "报销_приказ.docx")
    out_a = os.path.join(out_dir, "АктОтчет.xlsx")

    shutil.copy2(TMPL_DIR / "Смета_шаблон.docx", out_s)
    shutil.copy2(TMPL_DIR / "приказ_шаблон.docx", out_p)
    wb_tmpl = load_workbook(TMPL_DIR / "АктОтчет_шаблон.xlsx")
    wb_tmpl.save(out_a)

    doc = Document(out_s)
    set_cell_text(doc.tables[0].rows[0].cells[1], date_str)
    set_cell_text(doc.tables[1].rows[1].cells[1], str(budget_amt))
    set_cell_text(doc.tables[1].rows[2].cells[1], str(budget_amt))
    doc.save(out_s)

    doc = Document(out_p)
    DATE_RE = re.compile(r"(\d{2})\.(\d{2})\.(\d{4})")
    for para in doc.paragraphs:
        full = get_para_full(para)
        if "Красногорск" in full and DATE_RE.search(full):
            replace_para_text(para, DATE_RE.sub(date_str + " г.", full))
            continue
        if DATE_RE.match(full.strip().rstrip("г.").strip()):
            replace_para_text(para, date_str + " г.")
            continue
        if "В целях установления" in full and "делегация" in full:
            idx = full.find("делегация")
            replace_para_text(para, full[:idx] + "делегация  " + company_full)
            continue
        if "Провести" in full and "провести официальный прием" in full:
            new_text = re.sub(r"Провести\s+\d{1,2}\s+\w+\s+\d{4}\s+г\.",
                            "Провести " + date_russian, full)
            replace_para_text(para, new_text)
            continue
        if "смету представительских расходов" in full and "размере" in full:
            new_text = re.sub(r"размере\s+.+?\s+рублей\s+\d{2}\s+копеек",
                            "размере " + actual_words + " рублей 00 копеек", full)
            replace_para_text(para, new_text)
            continue
        if "Ответственному за представительское мероприятие работнику" in full:
            idx = full.find("работнику")
            replace_para_text(para, full[:idx] + "работнику " + responsible)
            continue
        if full.strip().startswith("-") and ("менеджер" in full.lower() or "Региональный" in full):
            replace_para_text(para, "- " + position + "  " + participant)
            continue
        if "/________________" in full and "(" in full:
            new_text = re.sub(r"\([^)]*\)", "(" + participant + ")", full)
            replace_para_text(para, new_text)
            continue
    doc.save(out_p)

    wb = load_workbook(out_a)
    ws = wb.active
    ws["D11"] = int(day); ws["G11"] = month; ws["J11"] = int(year)
    ws["B18"] = company_full
    ws["M22"] = int(day); ws["P22"] = month; ws["S22"] = int(year)
    ws["K23"] = venue; ws["K24"] = address
    for r in range(28, 74):
        ws["B" + str(r)] = None; ws["Y" + str(r)] = None
    for i in range(18):
        pk = "deleg_pos_" + str(i)
        nk = "deleg_nam_" + str(i)
        if pk in data and nk in data:
            pos = data[pk].strip()
            nam = data[nk].strip()
            if pos and nam:
                row = 28 + i * 2
                ws["B" + str(row)] = pos
                ws["Y" + str(row)] = nam
                ws["B" + str(row+1)] = "(должность)"
                ws["Y" + str(row+1)] = "(ФИО)"
    ws["B77"] = position; ws["Y77"] = participant
    ws["B81"] = purpose; ws["B84"] = result
    ws["N89"] = actual_amt; ws["S89"] = actual_words; ws["AE89"] = 0
    ws["Q94"] = rec_date_clean
    ws["U94"] = int(receipt_num) if str(receipt_num).isdigit() else receipt_num
    ws["Y94"] = receipt_amt
    for r in range(99, 107):
        ws["I" + str(r)] = None; ws["AC" + str(r)] = None
    if comm1_pos: ws["I99"] = comm1_pos
    if comm1_name: ws["AC99"] = comm1_name
    if comm2_pos: ws["I101"] = comm2_pos
    if comm2_name: ws["AC101"] = comm2_name
    if comm3_pos: ws["I103"] = comm3_pos
    if comm3_name: ws["AC103"] = comm3_name
    if compiler_pos: ws["I105"] = compiler_pos
    if compiler_name: ws["AC105"] = compiler_name
    wb.save(out_a)
    return out_s, out_p, out_a

# ============================================================
# UI
# ============================================================
st.set_page_config(page_title="报销生成器", page_icon="📋", layout="wide")
st.title("📋 效率报销-值得拥有")
st.caption("填写表单 → 自动生成代表团 → 一键下载")

if "delegation" not in st.session_state:
    st.session_state.delegation = []

with st.form("form"):
    c1, c2 = st.columns(2)
    with c1:
        st.subheader("基本信息")
        date_str = st.text_input("活动日期 *", "25.05.2026")
        actual_amt = st.number_input("实际报销金额 *", 0, value=22370)
        budget_amt = st.number_input("Смета预算金额 *", 0, value=30000, help="比实际大且取整")
        amount_manual = st.text_input("金额大写(不填=自动)", "")
        st.subheader("公司类型")
        company_type = st.selectbox("对方公司 *", ["russian", "chinese"],
                                     format_func=lambda x: "🇷🇺 俄罗斯公司" if x=="russian" else "🇨🇳 中国公司 (俄语转写)")
        st.subheader("公司与人员")
        submitter = st.text_input("提交人 (您的名字)", "")
        responsible = st.text_input("负责人 *", "Иван Ли")
        position = st.text_input("参与人职位 *", "Менеджер По Продажам")
        participant = st.text_input("参与人姓名 *", "Иван Ли")

    with c2:
        st.subheader("会议地点")
        venue = st.text_input("地点名称 *", "чуаньюй")
        address = st.text_input("地址 *", "МОСКВА, ПРОСПЕКТ.МИЧУРИНСКИЙ, ДОМ 7")
        st.subheader("会议内容")
        purpose = st.text_input("目的", "Продаж наши машины")
        result = st.text_input("成果", "Судим по проектом")
        st.subheader("小票")
        receipt_date = st.text_input("小票日期", "2026-05-24")
        receipt_num = st.text_input("小票编号", "")
        receipt_amt = st.number_input("小票金额(不填=实际)", 0, value=0)

    # --- Delegation ---
    st.subheader("👥 对方代表团 (每人 2300 ₽)")
    delegate_count = max(1, (actual_amt + 2299) // 2300)
    st.info(f"💰 {actual_amt} ₽ ÷ 2300 = **{delegate_count} 人** / 预算 **{delegate_count * 2300}** ₽")

    # Buttons: Generate + Clear (outside form submit, but we use form_submit_button)
    # Actually we need these as regular buttons that trigger rerun
    # For Streamlit forms, we can use columns for multiple buttons

    # Use regular st.button outside the form for the clear action
    st.markdown("---")

    dpos, dnam = [], []
    if st.session_state.delegation:
        tag = "🇨🇳" if company_type=="chinese" else "🇷🇺"
        st.caption(f"{tag} 已生成 {len(st.session_state.delegation)} 人 — 可手动修改任意字段")
        dcols_h = st.columns([3, 3, 1])
        with dcols_h[0]: st.markdown("**职位**")
        with dcols_h[1]: st.markdown("**姓名**")
        with dcols_h[2]: st.markdown("**预算**")
        for i, (pos, name) in enumerate(st.session_state.delegation):
            dc1, dc2, dc3 = st.columns([3, 3, 1])
            with dc1:
                dpos.append(st.text_input("dpos"+str(i), pos, key="gdp"+str(i), label_visibility="collapsed"))
            with dc2:
                dnam.append(st.text_input("dnam"+str(i), name, key="gdn"+str(i), label_visibility="collapsed"))
            with dc3:
                st.markdown("2300 ₽")
    else:
        st.caption("点击下方按钮自动生成，或手动填写")
        for i in range(delegate_count):
            dc1, dc2 = st.columns(2)
            with dc1: dpos.append(st.text_input("职位"+str(i+1), key="mdp"+str(i)))
            with dc2: dnam.append(st.text_input("姓名"+str(i+1), key="mdn"+str(i)))
        for i in range(delegate_count, 18):
            dpos.append(""); dnam.append("")

    while len(dpos) < 18: dpos.append("")
    while len(dnam) < 18: dnam.append("")

    # Action buttons row
    bcol1, bcol2, bcol3, bcol4 = st.columns([1, 1, 1, 2])
    with bcol1:
        gen_btn = st.form_submit_button("🎲 生成名单", type="secondary")
    with bcol2:
        clear_btn = st.form_submit_button("🗑️ 清空名单", type="secondary")
    with bcol3:
        recalc_btn = st.form_submit_button("🔄 重新计算", type="secondary")

    # --- Commission ---
    st.subheader("👤 Комиссия 成员")
    cm1, cm2 = st.columns(2)
    with cm1:
        comm1_pos = st.text_input("Ген.директор 职位", "Генеральный директор", key="c1p", disabled=True)
        comm1_name = st.text_input("Ген.директор 姓名", "Сяо Юаньсян", key="c1n", disabled=True)
    with cm2:
        comm2_pos = st.text_input("上级领导 职位(I101)", "", key="c2p")
        comm2_name = st.text_input("上级领导 姓名(AC101)", "", key="c2n")
    cm3, cm4 = st.columns(2)
    with cm3:
        comm3_pos = st.text_input("主会计 职位(I103)", "Главный бухгалтер", key="c3p")
        comm3_name = st.text_input("主会计 姓名(AC103)", "", key="c3n")
    with cm4:
        compiler_pos_ui = st.text_input("编制人 职位(I105)", "менеджер по продажам", key="cmp_p")
        compiler_name_ui = st.text_input("编制人 姓名(AC105)", participant, key="cmp_n")

    company_short = st.text_input("对方公司简称(不填=全名)", "")

    go = st.form_submit_button("🚀 生成报销文件", type="primary", use_container_width=True)

# --- Session state for usage log ---
if "usage_log" not in st.session_state:
    st.session_state.usage_log = []

# --- Handle delegation buttons ---
if "recalc_btn" in locals() and recalc_btn:
    st.session_state.delegation = generate_delegation(company_type, delegate_count)
    st.rerun()
if "gen_btn" in locals() and gen_btn:
    st.session_state.delegation = generate_delegation(company_type, delegate_count)
    st.rerun()

if "clear_btn" in locals() and clear_btn:
    st.session_state.delegation = []
    st.rerun()

# --- Usage log display ---
if st.session_state.usage_log:
    with st.expander("📊 使用记录 (本次会话)", expanded=False):
        for i, entry in enumerate(reversed(st.session_state.usage_log[-10:])):
            st.caption(f"{entry['time']} | {entry['submitter']} | {entry['company']} | {entry['actual_amt']}₽ (预算{entry['budget_amt']}₽) | {entry['delegates']}人 | 负责人:{entry['responsible']}")

if go:
    data = {
        "date": date_str, "actual_amt": actual_amt, "budget_amt": budget_amt,
        "amount_manual": amount_manual, "company_full": company_full,
        "responsible": responsible, "position": position, "participant": participant,
        "venue": venue, "address": address, "purpose": purpose, "result": result,
        "receipt_date": receipt_date, "receipt_num": receipt_num,
        "receipt_amt": receipt_amt if receipt_amt else actual_amt,
        "company_short": company_short,
        "comm1_pos": comm1_pos, "comm1_name": comm1_name,
        "comm2_pos": comm2_pos, "comm2_name": comm2_name,
        "comm3_pos": comm3_pos, "comm3_name": comm3_name,
        "compiler_pos": compiler_pos_ui, "compiler_name": compiler_name_ui,
    }
    for i,(p,n) in enumerate(zip(dpos,dnam)):
        data["deleg_pos_"+str(i)] = p
        data["deleg_nam_"+str(i)] = n

    required = {"活动日期": date_str, "实际报销金额": actual_amt, "对方公司全名": company_full,
                "负责人": responsible, "参与人职位": position, "参与人姓名": participant}
    missing = [k for k,v in required.items() if not v]
    if missing:
        st.error("请填写: " + ", ".join(missing))
    else:
        with st.spinner("正在生成..."):
            td = tempfile.mkdtemp()
            try:
                s,p,a = generate_files(data, td)
                                # Log the submission
                log_entry = {
                    "time": str(datetime.datetime.now().strftime("%Y-%m-%d %H:%M")),
                    "submitter": submitter if "submitter" in dir() else "",
                    "company": company_full,
                    "actual_amt": actual_amt,
                    "budget_amt": budget_amt,
                    "delegates": delegate_count,
                    "responsible": responsible,
                }
                st.session_state.usage_log.append(log_entry)
                st.success("生成完成! 预算:" + str(budget_amt) + " RUB | 实际:" + str(actual_amt) + " RUB")
                with open(s,"rb") as f: sb = f.read()
                with open(p,"rb") as f: pb = f.read()
                with open(a,"rb") as f: ab = f.read()
                zb = io.BytesIO()
                with zipfile.ZipFile(zb,"w") as zf:
                    zf.writestr("Смета.docx", sb)
                    zf.writestr("报销_приказ.docx", pb)
                    zf.writestr("АктОтчет.xlsx", ab)
                zb.seek(0)
                st.download_button("📥 下载全部 (ZIP)", zb, "报销文件.zip", "application/zip", use_container_width=True)
                cld, crd = st.columns(2)
                with cld: st.download_button("📄 Смета.docx", sb, "Смета.docx")
                with crd: st.download_button("📄 报销_приказ.docx", pb, "报销_приказ.docx")
                st.download_button("📊 АктОтчет.xlsx", ab, "АктОтчет.xlsx")
            finally:
                shutil.rmtree(td, ignore_errors=True)
